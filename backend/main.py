from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, EmailStr
from typing import List, Optional, Dict, Any
from datetime import datetime
import uvicorn
import logging
import os
from pathlib import Path
from email_service import email_service
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from fastapi.responses import Response

# Load environment variables from .env file
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="MindCare API", 
    description="API for MindCare Mental Health AI Counselor",
    version="1.0.0"
)

# CORS middleware - Allow all origins for development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with specific origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Groq client
try:
    from groq import Groq
    groq_api_key = os.getenv("GROQ_API_KEY")
    if groq_api_key:
        groq_client = Groq(api_key=groq_api_key)
        logger.info("Groq client initialized successfully")
    else:
        logger.warning("GROQ_API_KEY not found in environment variables")
        groq_client = None
except ImportError:
    logger.warning("Groq library not installed. Install with: pip install groq")
    groq_client = None
except Exception as e:
    logger.error(f"Failed to initialize Groq client: {e}")
    groq_client = None

# Pydantic models
class UserInfo(BaseModel):
    firstName: str
    lastName: str
    email: EmailStr
    phone: str
    age: int
    gender: str
    department: str
    reportTo: str  # "hr" or "manager"

class AssessmentData(BaseModel):
    answers: List[int]
    userInfo: UserInfo

class ChatMessage(BaseModel):
    message: str
    assessmentResults: Optional[Dict[str, Any]] = None
    sessionHistory: Optional[List[Dict[str, Any]]] = None

class ContactMessage(BaseModel):
    name: str
    email: EmailStr
    message: str

class SessionEndData(BaseModel):
    assessmentResults: Dict[str, Any]
    chatHistory: List[Dict[str, Any]]
    userInfo: UserInfo
    recommendations: List[Dict[str, str]]

# Helper functions
def calculate_dass_scores(answers: List[int]) -> Dict[str, Any]:
    """Calculate DASS-21 scores from answers"""
    # DASS-21 scoring: Depression (3,5,10,13,16,17,21), Anxiety (2,4,7,9,15,19,20), Stress (1,6,8,11,12,14,18)
    depression_items = [2, 4, 9, 12, 15, 16, 20]  # 0-indexed
    anxiety_items = [1, 3, 6, 8, 14, 18, 19]
    stress_items = [0, 5, 7, 10, 11, 13, 17]

    depression = sum(answers[i] for i in depression_items if i < len(answers)) * 2
    anxiety = sum(answers[i] for i in anxiety_items if i < len(answers)) * 2
    stress = sum(answers[i] for i in stress_items if i < len(answers)) * 2

    return {
        "depression": {
            "score": depression,
            "level": get_score_level("depression", depression),
        },
        "anxiety": {
            "score": anxiety,
            "level": get_score_level("anxiety", anxiety),
        },
        "stress": {
            "score": stress,
            "level": get_score_level("stress", stress),
        },
    }

def get_score_level(score_type: str, score: int) -> str:
    """Get severity level for DASS-21 scores"""
    ranges = {
        "depression": {
            "normal": [0, 9],
            "mild": [10, 13],
            "moderate": [14, 20],
            "severe": [21, 27],
            "extremely_severe": [28, 42],
        },
        "anxiety": {
            "normal": [0, 7],
            "mild": [8, 9],
            "moderate": [10, 14],
            "severe": [15, 19],
            "extremely_severe": [20, 42],
        },
        "stress": {
            "normal": [0, 14],
            "mild": [15, 18],
            "moderate": [19, 25],
            "severe": [26, 33],
            "extremely_severe": [34, 42],
        },
    }

    type_ranges = ranges[score_type]

    if score >= type_ranges["extremely_severe"][0]:
        return "Extremely Severe"
    if score >= type_ranges["severe"][0]:
        return "Severe"
    if score >= type_ranges["moderate"][0]:
        return "Moderate"
    if score >= type_ranges["mild"][0]:
        return "Mild"
    return "Normal"

def check_severe_case(assessment_results: Dict[str, Any]) -> bool:
    """Check if assessment indicates a severe case"""
    return any(
        result["level"] in ["Severe", "Extremely Severe"]
        for result in assessment_results.values()
    )

async def generate_groq_response(message: str, assessment_results: Optional[Dict[str, Any]] = None, session_history: Optional[List[Dict[str, Any]]] = None) -> str:   
    """Generate AI response using Groq API"""
    if not groq_client:
        return "I'm sorry, the AI service is currently unavailable. Please try again later."
    
    try:
        # Create context based on assessment results
        context = "You are a compassionate AI mental health counselor. Provide supportive, empathetic responses."
        
       
        # Create context based on assessment results
        context = "You are a compassionate AI mental health counselor. Provide supportive, empathetic responses."

        if assessment_results:
            depression_level = assessment_results.get("depression", {}).get("level", "Normal")
            anxiety_level = assessment_results.get("anxiety", {}).get("level", "Normal")
            stress_level = assessment_results.get("stress", {}).get("level", "Normal")

            context += f" The user's assessment shows: Depression: {depression_level}, Anxiety: {anxiety_level}, Stress: {stress_level}. Tailor your response accordingly."

        # Now, the 'messages_for_groq' part that follows will correctly find 'system_context'
        messages_for_groq = [{"role": "system", "content": context}] # This line should now work
        
        # Generate response using Groq
        

        if session_history:
                messages_for_groq.extend(session_history)

        messages_for_groq.append({"role": "user", "content": message})

        chat_completion = groq_client.chat.completions.create(
                messages=messages_for_groq, # Changed from 'messages' to 'messages_for_groq'
                model="llama3-8b-8192",  # or "mixtral-8x7b-32768"
                temperature=0.7,
                max_tokens=500
            )
               
        
        return chat_completion.choices[0].message.content
        
    except Exception as e:
        logger.error(f"Error generating Groq response: {e}")
        return "I'm having trouble processing your message right now. Could you please try rephrasing that?"

# Setup static files serving BEFORE defining routes
def setup_static_files():
    """Setup static file serving for frontend"""
    try:
        # Get the current directory (backend/)
        current_dir = Path(__file__).parent
        # Go up one level to project root
        project_root = current_dir.parent
        
        logger.info(f"Current directory: {current_dir}")
        logger.info(f"Project root: {project_root}")
        
        # Check if frontend files exist in project root
        frontend_files = ["index.html", "styles.css", "script.js", "counseling.html", "about.html"]
        missing_files = []
        
        for file in frontend_files:
            file_path = project_root / file
            if not file_path.exists():
                missing_files.append(file)
            else:
                logger.info(f"Found frontend file: {file_path}")
        
        if missing_files:
            logger.warning(f"Missing frontend files: {missing_files}")
            return False
        
        # Mount static files
        app.mount("/static", StaticFiles(directory=str(project_root)), name="static")
        logger.info(f"Static files mounted from: {project_root}")
        return True
        
    except Exception as e:
        logger.error(f"Error setting up static files: {e}")
        return False

# Setup static files
static_setup_success = setup_static_files()

# API Routes
@app.get("/")
async def root():
    """Serve the main index.html file"""
    try:
        current_dir = Path(__file__).parent
        project_root = current_dir.parent
        index_path = project_root / "index.html"
        
        if index_path.exists():
            return FileResponse(str(index_path))
        else:
            return {
                "message": "MindCare AI Counselor API",
                "status": "running",
                "version": "1.0.0",
                "groq_status": "connected" if groq_client else "disconnected",
                "static_files": "configured" if static_setup_success else "not found",
                "timestamp": datetime.now().isoformat(),
                "note": "Frontend files not found. Please ensure HTML files are in the project root."
            }
    except Exception as e:
        logger.error(f"Error serving index.html: {e}")
        return {
            "message": "MindCare AI Counselor API",
            "status": "running",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.get("/health")
async def health_check():
    """Detailed health check"""
    return {
        "status": "healthy",
        "service": "MindCare API",
        "timestamp": datetime.now().isoformat(),
        "groq_api": "connected" if groq_client else "disconnected",
        "static_files": "configured" if static_setup_success else "not configured",
        "endpoints": {
            "assessment": "/api/assess",
            "chat": "/api/chat",
            "contact": "/api/contact",
            "end_session": "/api/end-session",
            "documentation": "/docs"
        }
    }

# Serve specific HTML files
@app.get("/counseling.html")
async def serve_counseling():
    """Serve counseling.html"""
    try:
        current_dir = Path(__file__).parent
        project_root = current_dir.parent
        counseling_path = project_root / "counseling.html"
        
        if counseling_path.exists():
            return FileResponse(str(counseling_path))
        else:
            raise HTTPException(status_code=404, detail="Counseling page not found")
    except Exception as e:
        logger.error(f"Error serving counseling.html: {e}")
        raise HTTPException(status_code=500, detail="Error serving counseling page")

@app.get("/about.html")
async def serve_about():
    """Serve about.html"""
    try:
        current_dir = Path(__file__).parent
        project_root = current_dir.parent
        about_path = project_root / "about.html"
        
        if about_path.exists():
            return FileResponse(str(about_path))
        else:
            raise HTTPException(status_code=404, detail="About page not found")
    except Exception as e:
        logger.error(f"Error serving about.html: {e}")
        raise HTTPException(status_code=500, detail="Error serving about page")

# Serve CSS and JS files
@app.get("/styles.css")
async def serve_styles():
    """Serve styles.css"""
    try:
        current_dir = Path(__file__).parent
        project_root = current_dir.parent
        styles_path = project_root / "styles.css"
        
        if styles_path.exists():
            return FileResponse(str(styles_path), media_type="text/css")
        else:
            raise HTTPException(status_code=404, detail="Styles not found")
    except Exception as e:
        logger.error(f"Error serving styles.css: {e}")
        raise HTTPException(status_code=500, detail="Error serving styles")

@app.get("/script.js")
async def serve_script():
    """Serve script.js"""
    try:
        current_dir = Path(__file__).parent
        project_root = current_dir.parent
        script_path = project_root / "script.js"
        
        if script_path.exists():
            return FileResponse(str(script_path), media_type="application/javascript")
        else:
            raise HTTPException(status_code=404, detail="Script not found")
    except Exception as e:
        logger.error(f"Error serving script.js: {e}")
        raise HTTPException(status_code=500, detail="Error serving script")

@app.get("/counseling.js")
async def serve_counseling_script():
    """Serve counseling.js"""
    try:
        current_dir = Path(__file__).parent
        project_root = current_dir.parent
        script_path = project_root / "counseling.js"
        
        if script_path.exists():
            return FileResponse(str(script_path), media_type="application/javascript")
        else:
            raise HTTPException(status_code=404, detail="Counseling script not found")
    except Exception as e:
        logger.error(f"Error serving counseling.js: {e}")
        raise HTTPException(status_code=500, detail="Error serving counseling script")

@app.get("/auth.js")
async def serve_auth_script():
    """Serve auth.js"""
    try:
        current_dir = Path(__file__).parent
        project_root = current_dir.parent
        script_path = project_root / "auth.js"
        
        if script_path.exists():
            return FileResponse(str(script_path), media_type="application/javascript")
        else:
            raise HTTPException(status_code=404, detail="Auth script not found")
    except Exception as e:
        logger.error(f"Error serving auth.js: {e}")
        raise HTTPException(status_code=500, detail="Error serving auth script")

@app.get("/index.html")
async def serve_index_html():
    """Serve index.html explicitly"""
    try:
        current_dir = Path(__file__).parent
        project_root = current_dir.parent
        index_path = project_root / "index.html"
        
        if index_path.exists():
            return FileResponse(str(index_path))
        else:
            raise HTTPException(status_code=404, detail="Index page not found")
    except Exception as e:
        logger.error(f"Error serving index.html: {e}")
        raise HTTPException(status_code=500, detail="Error serving index page")

@app.post("/api/assess")
async def assess_mental_health(data: AssessmentData):
    """Process DASS-21 assessment and return results"""
    try:
        logger.info(f"Processing assessment for user: {data.userInfo.email}")
        logger.info(f"Received {len(data.answers)} answers: {data.answers}")
        
        # Validate answers
        if len(data.answers) != 21:
            logger.error(f"Invalid number of answers: {len(data.answers)}, expected 21")
            raise HTTPException(status_code=400, detail=f"Expected 21 answers, received {len(data.answers)}")
        
        # Validate answer values (should be 0-3)
        invalid_answers = [i for i, answer in enumerate(data.answers) if answer not in [0, 1, 2, 3]]
        if invalid_answers:
            logger.error(f"Invalid answer values at positions: {invalid_answers}")
            raise HTTPException(status_code=400, detail=f"Invalid answer values at positions: {invalid_answers}")
        
        # Calculate DASS-21 scores
        results = calculate_dass_scores(data.answers)
        logger.info(f"Calculated scores: {results}")
        
        # Check for severe cases
        is_severe = check_severe_case(results)
        
        if is_severe:
            logger.warning(f"Severe case detected for user: {data.userInfo.email}")
            # Send email notification in background
            try:
                await send_severe_case_notification(data.userInfo, results)
                logger.info("Severe case notification sent successfully")
            except Exception as email_error:
                logger.error(f"Failed to send severe case notification: {email_error}")
                # Don't fail the assessment if email fails
        
        # Log assessment completion
        logger.info(f"Assessment completed successfully - Depression: {results['depression']['score']}, "
                   f"Anxiety: {results['anxiety']['score']}, Stress: {results['stress']['score']}")
        
        response_data = {
            "success": True,
            "results": results,
            "userInfo": data.userInfo.dict(),
            "severeCaseDetected": is_severe,
            "timestamp": datetime.now().isoformat()
        }
        
        logger.info("Sending assessment response")
        return response_data
        
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        logger.error(f"Unexpected error processing assessment: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to process assessment: {str(e)}")

@app.post("/api/chat")
async def chat_with_counselor(data: ChatMessage):
    """Process chat messages with AI counselor using Groq"""
    try:
        logger.info(f"Processing chat message: {data.message[:50]}...")
        
        # Use Groq API for response generation
        if groq_client:
            response = await generate_groq_response(data.message, data.assessmentResults, data.sessionHistory)
            should_end = False
            recommendations = None
            
            # Check if user wants to end session
            end_keywords = ["end session", "finish", "goodbye", "end chat", "stop", "bye", "quit"]
            if any(keyword in data.message.lower() for keyword in end_keywords):
                should_end = True
                recommendations = generate_recommendations_from_assessment(data.assessmentResults)
                response += " Thank you for our session today. I've prepared some recommendations for you."
        
            # Fallback to local AI model
            
        
        logger.info(f"AI response generated, should_end: {should_end}")
        
        return {
            "success": True,
            "response": response,
            "shouldEndSession": should_end,
            "recommendations": recommendations if should_end else None,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error processing chat message: {str(e)}")
        # Return a fallback response instead of failing
        return {
            "success": True,
            "response": "I'm having some technical difficulties right now. Could you please rephrase your message or try again?",
            "shouldEndSession": False,
            "recommendations": None,
            "timestamp": datetime.now().isoformat()
        }

@app.post("/api/end-session")
async def end_session(data: SessionEndData):
    """Handle session end and send email report"""
    try:
        logger.info(f"Ending session for user: {data.userInfo.email}")
        
        # Check if severe case and send email
        is_severe = check_severe_case(data.assessmentResults)
        
        if is_severe:
            await send_severe_case_notification(data.userInfo, data.assessmentResults)
        
        # Send session report email
        await send_session_report(data.userInfo, data.assessmentResults, data.recommendations, data.chatHistory)
        
        return {
            "success": True,
            "message": "Session ended successfully. Report sent to designated recipient.",
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error ending session: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to end session")

@app.post("/api/contact")
async def contact_form(data: ContactMessage):
    """Handle contact form submissions"""
    try:
        logger.info(f"Contact form submission from: {data.email}")
        
        # In production, you would send an actual email
        logger.info(f"Contact message from {data.name} ({data.email}): {data.message}")
        
        return {
            "success": True,
            "message": "Thank you for your message. We'll get back to you soon!",
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error processing contact form: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to send message")

@app.get("/api/recommendations/{assessment_type}")
async def get_recommendations(assessment_type: str):
    """Get recommendations based on assessment type"""
    try:
        recommendations = {
            "depression": [
                {
                    "title": "Professional Counseling",
                    "description": "Consider speaking with a licensed mental health professional."
                },
                {
                    "title": "Daily Routine",
                    "description": "Establish consistent daily routines for better mental health."
                }
            ],
            "anxiety": [
                {
                    "title": "Breathing Exercises",
                    "description": "Practice deep breathing and mindfulness techniques."
                },
                {
                    "title": "Gradual Exposure",
                    "description": "Gradually face anxiety-provoking situations."
                }
            ],
            "stress": [
                {
                    "title": "Stress Management",
                    "description": "Learn stress management and relaxation techniques."
                },
                {
                    "title": "Physical Activity",
                    "description": "Regular exercise can significantly reduce stress."
                }
            ]
        }
        
        return {
            "success": True,
            "recommendations": recommendations.get(assessment_type, []),
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error getting recommendations: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to get recommendations")

@app.post("/api/download-report")
async def download_word_report(data: SessionEndData):
    """Generate and download Word document report"""
    try:
        logger.info(f"Generating Word report for user: {data.userInfo.email}")
        
        # Generate Word document
        doc_bytes = generate_word_report(
            data.userInfo, 
            data.assessmentResults, 
            data.recommendations, 
            data.chatHistory
        )
        
        # Create filename
        filename = f"MindCare_Report_{data.userInfo.firstName}_{data.userInfo.lastName}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        # Return file response
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logger.error(f"Error generating Word report: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate report")

# Email functions
async def send_severe_case_notification(user_info: UserInfo, assessment_results: Dict[str, Any]):
    """Send email notification for severe cases"""
    try:
        # Determine recipient based on user preference
        recipient_email = get_recipient_email(user_info.reportTo, user_info.department)
        
        subject = f"URGENT: Mental Health Alert - {user_info.firstName} {user_info.lastName} ({user_info.department})"
        
        body = f"""
        <h2>URGENT: Mental Health Alert</h2>
        <p>This is an automated alert from the MindCare AI Counseling System.</p>
        
        <h3>Employee Information:</h3>
        <ul>
            <li><strong>Name:</strong> {user_info.firstName} {user_info.lastName}</li>
            <li><strong>Email:</strong> {user_info.email}</li>
            <li><strong>Department:</strong> {user_info.department}</li>
            <li><strong>Phone:</strong> {user_info.phone}</li>
        </ul>
        
        <h3>Assessment Results:</h3>
        <ul>
            <li><strong>Depression:</strong> {assessment_results['depression']['score']} ({assessment_results['depression']['level']})</li>
            <li><strong>Anxiety:</strong> {assessment_results['anxiety']['score']} ({assessment_results['anxiety']['level']})</li>
            <li><strong>Stress:</strong> {assessment_results['stress']['score']} ({assessment_results['stress']['level']})</li>
        </ul>
        
        <p><strong>Recommended Action:</strong> Please reach out to this employee as soon as possible to provide appropriate support and resources.</p>
        
        <p><em>This is an automated message from the MindCare AI Counseling System.</em></p>
        """
        
        await email_service.send_email(recipient_email, subject, body)
        logger.info(f"Severe case notification sent to {recipient_email}")
        
    except Exception as e:
        logger.error(f"Failed to send severe case notification: {e}")

async def send_session_report(user_info: UserInfo, assessment_results: Dict[str, Any], 
                             recommendations: List[Dict[str, str]], chat_history: List[Dict[str, Any]]):
    """Send session report email"""
    try:
        recipient_email = get_recipient_email(user_info.reportTo, user_info.department)
        
        subject = f"Mental Health Session Report - {user_info.firstName} {user_info.lastName} ({user_info.department})"
        
        # Create recommendations HTML
        recommendations_html = ""
        for rec in recommendations:
            recommendations_html += f"<li><strong>{rec['title']}:</strong> {rec['description']}</li>"
        
        body = f"""
        <h2>Mental Health Session Report</h2>
        
        <h3>Employee Information:</h3>
        <ul>
            <li><strong>Name:</strong> {user_info.firstName} {user_info.lastName}</li>
            <li><strong>Email:</strong> {user_info.email}</li>
            <li><strong>Department:</strong> {user_info.department}</li>
            <li><strong>Session Date:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</li>
        </ul>
        
        <h3>Assessment Results:</h3>
        <ul>
            <li><strong>Depression:</strong> {assessment_results['depression']['score']} ({assessment_results['depression']['level']})</li>
            <li><strong>Anxiety:</strong> {assessment_results['anxiety']['score']} ({assessment_results['anxiety']['level']})</li>
            <li><strong>Stress:</strong> {assessment_results['stress']['score']} ({assessment_results['stress']['level']})</li>
        </ul>
        
        <h3>Recommendations:</h3>
        <ul>
            {recommendations_html}
        </ul>
        
        <p><em>This report is confidential and should be handled according to company privacy policies.</em></p>
        """
        
        await email_service.send_email(recipient_email, subject, body)
        logger.info(f"Session report sent to {recipient_email}")
        
    except Exception as e:
        logger.error(f"Failed to send session report: {e}")

def get_recipient_email(report_to: str, department: str) -> str:
    """Get recipient email based on user preference and department"""
    # Email configuration
    hr_emails = {
        "default": "itsharshprateek@gmail.com",
        "IT": "itsharshprateek@gmail.com",
        "HR": "itsharshprateek@gmail.com",
        "Finance": "itsharshprateek@gmail.com",
        "Marketing": "itsharshprateek@gmail.com"
    }
    
    manager_emails = {
        "default": "harshstring123@gmail.com",
        "IT": "harshstring123@gmail.com",
        "HR": "harshstring123@gmail.com", 
        "Finance": "harshstring123@gmail.com",
        "Marketing": "harshstring123@gmail.com"
    }
    
    if report_to.lower() == "hr":
        return hr_emails.get(department, hr_emails["default"])
    else:
        return manager_emails.get(department, manager_emails["default"])

def generate_recommendations_from_assessment(assessment_results: Dict[str, Any]) -> List[Dict[str, str]]:
    """Generate recommendations based on assessment results"""
    recommendations = []
    
    if assessment_results["depression"]["score"] >= 14:
        recommendations.append({
            "title": "Professional Counseling",
            "description": "Consider speaking with a licensed mental health professional for depression support."
        })
    
    if assessment_results["anxiety"]["score"] >= 10:
        recommendations.append({
            "title": "Anxiety Management",
            "description": "Practice breathing exercises and mindfulness techniques to manage anxiety."
        })
    
    if assessment_results["stress"]["score"] >= 19:
        recommendations.append({
            "title": "Stress Reduction",
            "description": "Implement stress management techniques and consider workload adjustments."
        })
    
    # Always add general recommendations
    recommendations.extend([
        {
            "title": "Regular Exercise",
            "description": "Engage in physical activity to improve overall mental health."
        },
        {
            "title": "Social Support",
            "description": "Maintain connections with colleagues, friends, and family."
        }
    ])
    
    return recommendations

def generate_word_report(user_info: UserInfo, assessment_results: Dict[str, Any], 
                        recommendations: List[Dict[str, str]], chat_history: List[Dict[str, Any]]) -> bytes:
    """Generate a Word document report"""
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading('MindCare Mental Health Assessment Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = doc.add_paragraph(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Personal Information Section
        doc.add_heading('Personal Information', level=1)
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Name', f'{user_info.firstName} {user_info.lastName}'),
            ('Email', user_info.email),
            ('Department', user_info.department),
            ('Age', str(user_info.age)),
            ('Gender', user_info.gender),
            ('Report Recipient', 'HR Department' if user_info.reportTo == 'hr' else 'Department Manager')
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
        
        doc.add_paragraph()
        
        # Assessment Results Section
        doc.add_heading('DASS-21 Assessment Results', level=1)
        
        results_table = doc.add_table(rows=4, cols=3)
        results_table.style = 'Table Grid'
        
        # Header row
        hdr_cells = results_table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Score'
        hdr_cells[2].text = 'Severity Level'
        
        # Data rows
        results_data = [
            ('Depression', assessment_results['depression']['score'], assessment_results['depression']['level']),
            ('Anxiety', assessment_results['anxiety']['score'], assessment_results['anxiety']['level']),
            ('Stress', assessment_results['stress']['score'], assessment_results['stress']['level'])
        ]
        
        for i, (category, score, level) in enumerate(results_data, 1):
            row_cells = results_table.rows[i].cells
            row_cells[0].text = category
            row_cells[1].text = str(score)
            row_cells[2].text = level
        
        doc.add_paragraph()
        
        # Score Interpretation
        doc.add_heading('Score Interpretation', level=2)
        doc.add_paragraph('The DASS-21 uses the following severity ranges:')
        
        interpretation_table = doc.add_table(rows=6, cols=4)
        interpretation_table.style = 'Table Grid'
        
        # Header
        interp_hdr = interpretation_table.rows[0].cells
        interp_hdr[0].text = 'Severity'
        interp_hdr[1].text = 'Depression'
        interp_hdr[2].text = 'Anxiety'
        interp_hdr[3].text = 'Stress'
        
        # Data
        interp_data = [
            ('Normal', '0-9', '0-7', '0-14'),
            ('Mild', '10-13', '8-9', '15-18'),
            ('Moderate', '14-20', '10-14', '19-25'),
            ('Severe', '21-27', '15-19', '26-33'),
            ('Extremely Severe', '28+', '20+', '34+')
        ]
        
        for i, (severity, dep, anx, stress) in enumerate(interp_data, 1):
            row = interpretation_table.rows[i].cells
            row[0].text = severity
            row[1].text = dep
            row[2].text = anx
            row[3].text = stress
        
        doc.add_paragraph()
        
        # Recommendations Section
        doc.add_heading('Personalized Recommendations', level=1)
        
        for i, rec in enumerate(recommendations, 1):
            doc.add_heading(f'{i}. {rec["title"]}', level=2)
            doc.add_paragraph(rec['description'])
        
        # Session Summary
        if chat_history:
            doc.add_page_break()
            doc.add_heading('Counseling Session Summary', level=1)
            
            user_messages = [msg for msg in chat_history if msg.get('sender') == 'user']
            bot_messages = [msg for msg in chat_history if msg.get('sender') == 'bot']
            
            doc.add_paragraph(f'Total Messages Exchanged: {len(chat_history)}')
            doc.add_paragraph(f'User Messages: {len(user_messages)}')
            doc.add_paragraph(f'AI Counselor Responses: {len(bot_messages)}')
            
            doc.add_heading('Key Discussion Points', level=2)
            
            # Extract key themes from user messages
            key_themes = []
            for msg in user_messages[:5]:  # First 5 user messages
                if len(msg.get('message', '')) > 20:
                    key_themes.append(msg['message'][:100] + '...' if len(msg['message']) > 100 else msg['message'])
            
            for theme in key_themes:
                p = doc.add_paragraph(theme)
                p.style = 'List Bullet'
        
        # Emergency Resources
        doc.add_page_break()
        doc.add_heading('Emergency Resources', level=1)
        
        emergency_info = [
            'If you are in immediate danger, call 911',
            'National Suicide Prevention Lifeline: 988 or 1-800-273-8255',
            'Crisis Text Line: Text HOME to 741741',
            'National Domestic Violence Hotline: 1-800-799-7233',
            'SAMHSA National Helpline: 1-800-662-4357'
        ]
        
        for info in emergency_info:
            p = doc.add_paragraph(info)
            p.style = 'List Bullet'
        
        # Disclaimer
        doc.add_paragraph()
        doc.add_heading('Important Disclaimer', level=2)
        disclaimer_text = (
            "This assessment and counseling session are provided by an AI system and should not replace "
            "professional mental health care. If you are experiencing severe symptoms or are in crisis, "
            "please contact a licensed mental health professional or emergency services immediately. "
            "This report is confidential and should be handled according to applicable privacy laws and "
            "company policies."
        )
        doc.add_paragraph(disclaimer_text)
        
        # Save to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
        
    except Exception as e:
        logger.error(f"Error generating Word report: {e}")
        raise

# Error handlers
@app.exception_handler(404)
async def not_found_handler(request: Request, exc: HTTPException):
    return JSONResponse(
        status_code=404,
        content={
            "error": "Not Found",
            "message": f"The endpoint {request.url.path} was not found",
            "timestamp": datetime.now().isoformat()
        }
    )

@app.exception_handler(500)
async def internal_error_handler(request: Request, exc: HTTPException):
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal Server Error",
            "message": "An unexpected error occurred",
            "timestamp": datetime.now().isoformat()
        }
    )

# Run the application
if __name__ == "__main__":
    logger.info("Starting MindCare API server...")
    logger.info(f"Static files setup: {'Success' if static_setup_success else 'Failed'}")
    uvicorn.run(
        "main:app", 
        host="0.0.0.0", 
        port=8000, 
        reload=True,
        log_level="info"
    )
